# -*- coding: utf-8 -*-
"""Assemble the course-design report .docx from section markdown files.

Layout follows 报告模板.doc: title / author / 摘要 / numbered sections / references.
Markdown subset supported: # ## ### headings, pipe tables, ![caption](image),
**bold** inline, numbered & dash lists, plain paragraphs, standalone formula lines.
"""
import os
import re
import sys

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ---------------------------------------------------------------- low-level XML helpers
def _border_el(edge, sz, val="single", color="000000"):
    e = OxmlElement("w:" + edge)
    e.set(qn("w:val"), val)
    if val != "nil":
        e.set(qn("w:sz"), str(sz))
        e.set(qn("w:color"), color)
    return e


def _cell_set_bottom(cell, sz=6, color="000000"):
    tcPr = cell._tc.get_or_add_tcPr()
    tb = tcPr.find(qn("w:tcBorders"))
    if tb is None:
        tb = OxmlElement("w:tcBorders"); tcPr.append(tb)
    old = tb.find(qn("w:bottom"))
    if old is not None:
        tb.remove(old)
    tb.append(_border_el("bottom", sz, color=color))


def _cell_shade(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), fill)
    tcPr.append(shd)


def three_line(table, header_fill="F2F2F2"):
    """Classic academic three-line table: thick top & bottom rules, thin rule
    under the header row, subtle header shading, no vertical/inner lines."""
    tblPr = table._tbl.tblPr
    old = tblPr.find(qn("w:tblBorders"))
    if old is not None:
        tblPr.remove(old)
    b = OxmlElement("w:tblBorders")
    b.append(_border_el("top", 14))
    b.append(_border_el("bottom", 14))
    for edge in ("left", "right", "insideH", "insideV"):
        b.append(_border_el(edge, 0, val="nil"))
    tblPr.append(b)
    for c in table.rows[0].cells:
        _cell_set_bottom(c, sz=8)
        if header_fill:
            _cell_shade(c, header_fill)


def add_toc(doc):
    """Insert an auto-updating Table of Contents field (levels 1-3)."""
    p = doc.add_paragraph()
    run = p.add_run()
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), 'TOC \\o "1-2" \\h \\z \\u')
    r = OxmlElement("w:r"); t = OxmlElement("w:t")
    t.text = "（在 Word 中右键此处选择“更新域”可生成目录）"
    r.append(t); fld.append(r)
    p._p.append(fld)
    return p


def add_page_number_footer(doc):
    """Centered 'PAGE' field in the footer; blank on the title page."""
    sec = doc.sections[0]
    sec.different_first_page_header_footer = True
    fp = sec.footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = sec.footer.paragraphs[0]
    # primary footer (non-first pages)
    foot = sec.footer
    par = foot.paragraphs[0]
    par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = par.add_run()
    for kind, txt in (("begin", None), (None, "PAGE"), ("end", None)):
        if kind:
            fc = OxmlElement("w:fldChar"); fc.set(qn("w:fldCharType"), kind); run._r.append(fc)
        else:
            it = OxmlElement("w:instrText"); it.set(qn("xml:space"), "preserve"); it.text = txt
            run._r.append(it)
    set_fonts(run, size=10.5)

PROJ = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
SEC = os.path.join(PROJ, "report", "sections2")
FIG = os.path.join(PROJ, "figures")

TITLE = "面向预训练视觉Transformer的参数高效微调:跨方法、跨规模、跨数据集的系统研究与改进"
AUTHOR = "胡昊铭，3024210028"
OUT = os.path.join(PROJ, "report", "实验报告_胡昊铭_3024210028.docx")

SECTION_FILES = [
    "00_摘要.md", "01_引言.md", "02_相关工作.md", "03_方法A.md", "04_方法B.md",
    "05_实验设置.md", "06_主结果.md", "07_缩放分析.md", "08_数据效率.md",
    "09_消融.md", "10_可解释性.md", "11_效率.md", "12_讨论.md",
    "13_结论.md", "14_参考文献.md", "15_附录.md",
]


# ---------------------------------------------------------------- helpers
def set_fonts(run, ascii_font="Times New Roman", east_font="宋体",
              size=10.5, bold=False, italic=False):
    run.font.name = ascii_font
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = rPr.makeelement(qn("w:rFonts"), {})
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), east_font)


BOLD_RE = re.compile(r"\*\*(.+?)\*\*")


def add_runs_with_bold(par, text, **font_kw):
    """Split **bold** spans into separate runs."""
    pos = 0
    for m in BOLD_RE.finditer(text):
        if m.start() > pos:
            set_fonts(par.add_run(text[pos:m.start()]), **font_kw)
        kw = dict(font_kw)
        kw["bold"] = True
        set_fonts(par.add_run(m.group(1)), **kw)
        pos = m.end()
    if pos < len(text):
        set_fonts(par.add_run(text[pos:]), **font_kw)


def body_par(doc, text, indent=True, east="宋体", size=12,
             align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    p = doc.add_paragraph()
    p.alignment = align
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = 1.5
    pf.space_after = Pt(4)
    if indent:
        pf.first_line_indent = Pt(size * 2)
    add_runs_with_bold(p, text, east_font=east, size=size)
    return p


_H1_SEEN = []


def heading_par(doc, text, level):
    # use Word heading styles so the TOC field can pick them up (outline level)
    style_name = {1: "Heading 1", 2: "Heading 2", 3: "Heading 3"}[level]
    p = doc.add_paragraph(style=style_name)
    pf = p.paragraph_format
    pf.keep_with_next = True
    if level == 1:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        pf.space_before = Pt(22) if _H1_SEEN else Pt(6)
        pf.space_after = Pt(10)
        _H1_SEEN.append(1)
        size = 15
    elif level == 2:
        pf.space_before, pf.space_after = Pt(13), Pt(6)
        size = 13.5
    else:
        pf.space_before, pf.space_after = Pt(10), Pt(5)
        size = 12
    add_runs_with_bold(p, text, east_font="黑体", size=size)
    for r in p.runs:
        r.font.bold = True
        r.font.color.rgb = RGBColor(0, 0, 0)
    return p


def caption_par(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(8)
    add_runs_with_bold(p, text, east_font="宋体", size=10.5, bold=True)
    for r in p.runs:
        r.font.bold = True
    return p


def add_image(doc, path, caption):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    run = p.add_run()
    run.add_picture(path, width=Cm(14.5))
    if caption:
        caption_par(doc, caption)


def add_table(doc, rows):
    """rows: list of list of cell strings; first row = header. Three-line style."""
    ncols = max(len(r) for r in rows)
    t = doc.add_table(rows=len(rows), cols=ncols)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, row in enumerate(rows):
        for j in range(ncols):
            cell = t.cell(i, j)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            text = row[j] if j < len(row) else ""
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(1.5)
            p.paragraph_format.space_after = Pt(1.5)
            p.paragraph_format.line_spacing = 1.0
            add_runs_with_bold(p, text, east_font="宋体", size=10)
            if i == 0:
                for r in p.runs:
                    r.font.bold = True
    three_line(t)
    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Pt(4)
    return t


IMG_RE = re.compile(r"^!\[(.*?)\]\((.*?)\)\s*$")
TABLE_ROW_RE = re.compile(r"^\s*\|(.+)\|\s*$")
TABLE_SEP_RE = re.compile(r"^\s*\|?\s*:?-{2,}.*$")
REF_RE = re.compile(r"^\[\d+\]\s")
NUMLIST_RE = re.compile(r"^\s*(\d+[\)\.])\s+")
CN_RE = re.compile(r"[一-鿿]")


def is_formula_line(line):
    """Standalone math-ish line: no Chinese, has = or math symbols, short."""
    if CN_RE.search(line):
        return False
    if len(line) > 90 or len(line) < 5:
        return False
    return bool(re.search(r"[=⊙←∈≈Δθγβα⋅·×]", line))


def parse_table_cells(line):
    inner = line.strip().strip("|")
    return [c.strip() for c in inner.split("|")]


def render_markdown(doc, md_text, skip_h1=False):
    lines = md_text.splitlines()
    i = 0
    while i < len(lines):
        raw = lines[i]
        line = raw.rstrip()
        s = line.strip()
        if not s:
            i += 1
            continue
        # code fence ``` ... ``` -> verbatim monospace, no markdown parsing inside
        if s.startswith("```"):
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code_line = lines[i].rstrip("\n").rstrip()
                cp = doc.add_paragraph()
                cp.paragraph_format.left_indent = Pt(14)
                cp.paragraph_format.space_after = Pt(0)
                cp.paragraph_format.space_before = Pt(0)
                cp.paragraph_format.line_spacing = 1.05
                r = cp.add_run(code_line if code_line else " ")
                r.font.name = "Consolas"
                r.font.size = Pt(9.5)
                i += 1
            i += 1  # skip closing fence
            continue
        # headings
        if s.startswith("### "):
            heading_par(doc, s[4:].strip(), 3); i += 1; continue
        if s.startswith("## "):
            heading_par(doc, s[3:].strip(), 2); i += 1; continue
        if s.startswith("# "):
            if not skip_h1:
                heading_par(doc, s[2:].strip(), 1)
            else:
                heading_par(doc, s[2:].strip(), 1)
            i += 1; continue
        # image
        m = IMG_RE.match(s)
        if m:
            cap, fname = m.group(1).strip(), os.path.basename(m.group(2).strip())
            path = os.path.join(FIG, fname)
            if os.path.exists(path):
                add_image(doc, path, cap)
            else:
                body_par(doc, f"[缺失图片: {fname}] {cap}", indent=False)
            i += 1; continue
        # table
        if TABLE_ROW_RE.match(s):
            rows = []
            while i < len(lines) and TABLE_ROW_RE.match(lines[i].strip()):
                if not TABLE_SEP_RE.match(lines[i].strip()):
                    rows.append(parse_table_cells(lines[i]))
                i += 1
            if rows:
                add_table(doc, rows)
            continue
        # reference entries
        if REF_RE.match(s):
            p = body_par(doc, s, indent=False, size=11)
            p.paragraph_format.left_indent = Pt(18)
            p.paragraph_format.first_line_indent = Pt(-18)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            i += 1; continue
        # table/figure caption written as bold standalone line
        if (s.startswith("**表") or s.startswith("**图")) and s.endswith("**"):
            caption_par(doc, s.strip("*")); i += 1; continue
        # formula line
        if is_formula_line(s):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(3)
            set_fonts(p.add_run(s), ascii_font="Times New Roman",
                      east_font="宋体", size=12, italic=False)
            i += 1; continue
        # numbered list item
        if NUMLIST_RE.match(s) or s.startswith("- "):
            text = s[2:].strip() if s.startswith("- ") else s
            p = body_par(doc, text, indent=False)
            p.paragraph_format.left_indent = Pt(21)
            i += 1; continue
        # plain paragraph
        body_par(doc, s)
        i += 1


def main():
    doc = Document()
    sec = doc.sections[0]
    sec.page_width, sec.page_height = Cm(21.0), Cm(29.7)  # A4
    for attr in ("left_margin", "right_margin"):
        setattr(sec, attr, Cm(2.5))
    for attr in ("top_margin", "bottom_margin"):
        setattr(sec, attr, Cm(2.5))

    # default style
    st = doc.styles["Normal"]
    st.font.name = "Times New Roman"
    st.font.size = Pt(12)
    st.element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

    # configure heading styles (black, 黑体) so TOC entries look clean
    for name, sz in (("Heading 1", 15), ("Heading 2", 13.5), ("Heading 3", 12)):
        hs = doc.styles[name]
        hs.font.name = "Times New Roman"
        hs.font.size = Pt(sz)
        hs.font.bold = True
        hs.font.color.rgb = RGBColor(0, 0, 0)
        hs.element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")

    add_page_number_footer(doc)

    # ---------------- title page (vertically centred block) ----------------
    for _ in range(6):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(18)
    p.paragraph_format.line_spacing = 1.4
    set_fonts(p.add_run(TITLE), east_font="黑体", size=20, bold=True)
    for line, sz in (("《机器学习与深度学习》课程设计", 14), (AUTHOR, 13), ("2026 年 6 月", 13)):
        q = doc.add_paragraph()
        q.alignment = WD_ALIGN_PARAGRAPH.CENTER
        q.paragraph_format.space_before = Pt(10)
        set_fonts(q.add_run(line), east_font="宋体", size=sz)
    doc.add_page_break()

    # ---------------- table of contents ----------------
    tp = doc.add_paragraph()
    tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tp.paragraph_format.space_after = Pt(10)
    set_fonts(tp.add_run("目  录"), east_font="黑体", size=16, bold=True)
    add_toc(doc)
    doc.add_page_break()

    for f in SECTION_FILES:
        path = os.path.join(SEC, f)
        if not os.path.exists(path):
            print("MISSING SECTION:", f)
            continue
        md = open(path, encoding="utf-8").read()
        render_markdown(doc, md)

    doc.save(OUT)
    print("SAVED:", OUT)

    # quick self-check
    d2 = Document(OUT)
    n_imgs = len(d2.inline_shapes)
    n_tables = len(d2.tables)
    n_pars = len(d2.paragraphs)
    n_chars = sum(len(p.text) for p in d2.paragraphs) + sum(
        len(c.text) for t in d2.tables for row in t.rows for c in row.cells)
    print(f"CHECK: images={n_imgs} tables={n_tables} paragraphs={n_pars} chars={n_chars}")


if __name__ == "__main__":
    main()
